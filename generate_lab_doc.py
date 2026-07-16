#!/usr/bin/env python3
"""Generate the Network Automation Bootcamp Word document (v2.0)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Helpers ──────────────────────────────────────────────────────────────────

def heading(text, level=1, color=(0x1F, 0x4E, 0x79)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFEFEF")
    p._p.get_or_add_pPr().append(shd)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

# ── Cover ─────────────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Network Automation Bootcamp")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lab Guide — Git · Ansible · REST API")
r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Version 2.0   |   {datetime.date.today().strftime('%B %Y')}   |   10–12 Students")
r.font.size = Pt(11)

doc.add_page_break()

# ── 1. Introduction ───────────────────────────────────────────────────────────

heading("1. Introduction")
doc.add_paragraph(
    "The Network Automation Bootcamp is a hands-on lab designed for 10–12 network engineers. "
    "It covers three core pillars of modern network automation: Git, Ansible, and REST APIs. "
    "Students work on shared lab devices (two routers and one switch), a local NetBox instance "
    "as the source of truth, and several free public APIs to explore real-world data — no "
    "accounts or credit cards required."
)

heading("1.1 Learning Objectives", 2, (0x2E, 0x74, 0xB5))
for obj in [
    "Track and roll back network configs using Git branches and merge requests.",
    "Write Ansible playbooks to configure routers and a switch at scale with Jinja2 templates.",
    "Consume the NetBox REST API to query and update the source of truth.",
    "Query free public APIs (ipctl.io, PeeringDB, ipinfo.io) to explore real BGP and IP data.",
    "Connect all three tools into a repeatable, auditable end-to-end change pipeline.",
]:
    bullet(obj)

heading("1.2 Student Setup", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph("Each student receives:")
for item in [
    "A workstation IP in 10.0.0.101–112 (Pod-01 to Pod-12)",
    "A personal Git branch:  student/pod-XX-yourname",
    "A personal NetBox API token for the lab duration",
    "SSH credentials for the shared lab devices",
]:
    bullet(item)

code(
"git clone http://10.0.0.60/bootcamp/Network-Automation-Bootcamp.git\n"
"cd Network-Automation-Bootcamp\n"
"git checkout -b student/pod-XX-yourname\n"
"pip install ansible requests netmiko napalm pynetbox"
)

heading("1.3 Prerequisites", 2, (0x2E, 0x74, 0xB5))
table(
    ["Requirement", "Version", "Purpose"],
    [
        ["Git",        "≥ 2.30",  "Source control for network configs"],
        ["Python",     "≥ 3.9",   "Scripting and API interaction"],
        ["Ansible",    "≥ 2.14",  "Configuration management engine"],
        ["requests",   "latest",  "HTTP client for REST/RESTCONF"],
        ["netmiko",    "latest",  "Multi-vendor SSH sessions"],
        ["napalm",     "latest",  "Device abstraction layer"],
        ["pynetbox",   "latest",  "NetBox API Python wrapper"],
    ],
    [1.8, 1.2, 3.5]
)

# ── 2. Lab Architecture ───────────────────────────────────────────────────────

heading("2. Lab Architecture & Topology")
doc.add_paragraph(
    "The lab spans three networks: the classroom management network (10.0.0.0/24) where "
    "students and the Git server live; the automation & device network (10.106.106.0/24) "
    "where the control node and the routers/switch live; NetBox on its own network "
    "(10.100.100.0/24); and the internet for free public API access."
)

code(
"                         INTERNET\n"
"                             │\n"
"          ───────────────────┼──────────────────────\n"
"            FREE PUBLIC APIs (no account needed)\n"
"          ───────────────────┼──────────────────────\n"
"          │                  │                      │\n"
"      ipinfo.io         ipctl.io API         PeeringDB\n"
"       (free)             (free)               (free)\n"
"          │                  │                      │\n"
"          └──────────────────┼──────────────────────┘\n"
"                             │ HTTPS\n"
"═════════════════════════════╪══════════════════════════\n"
"         MGMT NETWORK  10.0.0.0/24\n"
"═════════════════════════════╪══════════════════════════\n"
"                             │\n"
"                        GIT SERVER            NETBOX\n"
"                        10.0.0.60          10.100.100.25\n"
"                        (Gitea)              :8000 (IPAM)\n"
"                             │                    │\n"
"═════════════════════════════╪════════════════════╪══════\n"
"     AUTOMATION & DEVICE NETWORK  10.106.106.0/24\n"
"═════════════════════════════╪════════════════════╪══════\n"
"                             │                    │\n"
"                       CONTROL NODE ◄── REST API ──┘\n"
"                       10.106.106.60\n"
"                       Git/Ansible/Python\n"
"                             │\n"
"                             │ SSH / RESTCONF / Ansible\n"
"       ┌─────────────────────┼───────────────┐\n"
"       │                     │               │\n"
"  ROUTER-1              ROUTER-2        SWITCH-1\n"
"  10.106.106.61          10.106.106.62   10.106.106.63\n"
"  NX-OS                  NX-OS           IOS  VLANs\n"
"\n"
"═══════════════════════════════════════════════════\n"
"  STUDENT PODS  10.0.0.101 – 10.0.0.112  (×12)\n"
"═══════════════════════════════════════════════════"
)

heading("2.1 IP Address Plan", 2, (0x2E, 0x74, 0xB5))
table(
    ["Device / Host", "IP Address", "Role"],
    [
        ["Control Node",     "10.106.106.60", "Automation engine (shared)"],
        ["NetBox",           "10.100.100.25", "IPAM / DCIM + REST API"],
        ["Gitea (Git Srv)",  "10.0.0.60",     "Repository server"],
        ["Router-1",         "10.106.106.61", "NX-OS core router (virtual Nexus)"],
        ["Router-2",         "10.106.106.62", "NX-OS secondary router (virtual Nexus)"],
        ["Switch-1",         "10.106.106.63", "IOS access switch + VLANs"],
        ["Student Pod 01–12","10.0.0.101–112","Student workstations"],
    ],
    [2.0, 1.5, 3.0]
)

heading("2.2 Free Public API Reference", 2, (0x2E, 0x74, 0xB5))
table(
    ["Service", "Base URL", "Auth", "Best For"],
    [
        ["NetBox Demo",         "demo.netbox.dev/api/",              "Public demo token",   "IPAM API queries"],
        ["ipctl.io",            "api.ipctl.io/v1",                   "None (core routes)",  "ASN, prefix, BGP, RPKI data"],
        ["PeeringDB",           "peeringdb.com/api/",                "None (read)",         "IX and peering records"],
        ["ipinfo.io",           "ipinfo.io",                         "None (50 k/mo)",      "IP geolocation & ASN"],
        ["httpbin.org",         "httpbin.org",                       "None",                "HTTP method practice"],
    ],
    [1.8, 1.8, 1.4, 1.5]
)

# ── 3. Git ────────────────────────────────────────────────────────────────────

heading("3. Section 1 — Git for Network Automation")
doc.add_paragraph(
    "Git is the single source of truth for all device configurations. Every change is "
    "tracked, attributed, and reversible. Students work on personal branches and submit "
    "merge requests to main, mirroring real change-management workflows."
)

heading("3.1 Exercise 1 — Initialise a Config Repository", 2, (0x2E, 0x74, 0xB5))
code(
"mkdir net-configs && cd net-configs\n"
"git init\n"
"git config user.name  'Network Engineer'\n"
"git config user.email 'neteng@example.com'\n"
"cp ../configs/router1_baseline.cfg .\n"
"git add router1_baseline.cfg\n"
"git commit -m 'feat: add Router-1 baseline config'\n"
"git log --oneline"
)

heading("3.2 Exercise 2 — Branching for Change Management", 2, (0x2E, 0x74, 0xB5))
code(
"git checkout -b change/CR-1042-add-loopback\n"
"echo 'interface Loopback0\\n ip address 10.255.255.1 255.255.255.255' >> router1.cfg\n"
"git add router1.cfg\n"
"git commit -m 'feat(CR-1042): add Loopback0 to Router-1'\n"
"git diff main change/CR-1042-add-loopback\n"
"git checkout main\n"
"git merge --no-ff change/CR-1042-add-loopback"
)

heading("3.3 Exercise 3 — Configuration Rollback", 2, (0x2E, 0x74, 0xB5))
code(
"echo 'no ip routing' >> router1.cfg\n"
"git add router1.cfg && git commit -m 'ERROR: removed ip routing'\n"
"git log --oneline\n"
"git revert HEAD          # safe: creates an undo commit\n"
"git push origin main"
)

heading("3.4 Key Git Commands", 2, (0x2E, 0x74, 0xB5))
table(
    ["Command", "Purpose"],
    [
        ["git init",                    "Initialise a new repository"],
        ["git clone <url>",             "Clone a remote repository"],
        ["git add <file>",              "Stage changes"],
        ["git commit -m 'msg'",         "Commit staged changes"],
        ["git log --oneline",           "Compact commit history"],
        ["git checkout -b <branch>",    "Create and switch to new branch"],
        ["git merge --no-ff <branch>",  "Merge with a merge commit"],
        ["git revert HEAD",             "Undo last commit (safe)"],
        ["git tag v1.0",                "Tag a golden config release"],
    ],
    [3.0, 3.5]
)

# ── 4. Ansible ────────────────────────────────────────────────────────────────

heading("4. Section 2 — Ansible for Network Automation")
doc.add_paragraph(
    "Ansible connects to the lab routers and switch over SSH and applies declarative YAML "
    "playbooks. Students always run with --check --diff before applying, mirroring the "
    "peer-review gate used in production."
)

heading("4.1 Lab Devices", 2, (0x2E, 0x74, 0xB5))
table(
    ["Host", "IP", "OS", "Role"],
    [
        ["router1", "10.106.106.61",  "NX-OS", "Core router / OSPF (no enable needed)"],
        ["router2", "10.106.106.62",  "NX-OS",  "Secondary router (no enable needed)"],
        ["switch1", "10.106.106.63", "IOS",    "Access switch / VLANs (needs enable)"],
    ],
    [1.2, 1.2, 1.0, 3.0]
)

heading("4.2 Inventory File (hosts.ini)", 2, (0x2E, 0x74, 0xB5))
code(
"[routers]\n"
"router1 ansible_host=10.106.106.61\n"
"router2 ansible_host=10.106.106.62\n"
"\n"
"[switches]\n"
"switch1 ansible_host=10.106.106.63\n"
"\n"
"[network:children]\n"
"routers\n"
"switches\n"
"\n"
"[network:vars]\n"
"ansible_user=admin\n"
"ansible_password=Lab@1234\n"
"ansible_connection=network_cli\n"
"\n"
"# Routers are NX-OS and skip enable; switch is IOS and needs it\n"
"[routers:vars]\n"
"ansible_network_os=nxos\n"
"ansible_become=no\n"
"\n"
"[switches:vars]\n"
"ansible_network_os=ios\n"
"ansible_become=yes\n"
"ansible_become_method=enable"
)

heading("4.3 Playbook 1 — Gather Facts", 2, (0x2E, 0x74, 0xB5))
code("ansible-playbook -i inventory/hosts.ini playbooks/01-gather-facts.yml")

heading("4.4 Playbook 2 — Push Interface Config", 2, (0x2E, 0x74, 0xB5))
code(
"ansible-playbook -i inventory/hosts.ini playbooks/02-push-config.yml --check --diff\n"
"ansible-playbook -i inventory/hosts.ini playbooks/02-push-config.yml"
)

heading("4.5 Playbook 3 — Compliance Check", 2, (0x2E, 0x74, 0xB5))
code("ansible-playbook -i inventory/hosts.ini playbooks/03-compliance-check.yml")

heading("4.6 Jinja2 Template Snippet", 2, (0x2E, 0x74, 0xB5))
code(
"{% for iface in interfaces %}\n"
"interface {{ iface.name }}\n"
" description {{ iface.description }}\n"
" {% if iface.ip is defined %}\n"
" ip address {{ iface.ip }} {{ iface.mask }}\n"
" {% endif %}\n"
" {% if iface.shutdown %}shutdown{% else %}no shutdown{% endif %}\n"
"!\n"
"{% endfor %}"
)

# ── 5. API ────────────────────────────────────────────────────────────────────

heading("5. Section 3 — REST API for Network Automation")
doc.add_paragraph(
    "This section uses two types of APIs: the local NetBox REST API as a source of truth, "
    "and several free, no-account public APIs for exploring real-world routing and network data."
)

heading("5.1 Script 1 — NetBox: Query Devices", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph(
    "Fetches active devices from the local NetBox instance and prints a summary table. "
    "Students then extend the script to POST a new device via the API."
)
code(
"NETBOX_URL   = 'https://10.100.100.25'\n"
"NETBOX_TOKEN = 'your_token_here'\n"
"\n"
"headers = {'Authorization': f'Token {NETBOX_TOKEN}'}\n"
"resp = requests.get(f'{NETBOX_URL}/api/dcim/devices/?status=active', headers=headers)\n"
"devices = resp.json()['results']"
)

heading("5.2 Script 2 — NetBox: Parsing API Output", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph(
    "Pages through every device instead of one page, then runs the results through several "
    "parsing patterns: grouping by site, counting by status, filtering to devices missing a "
    "primary IP, filtering by role, and flattening to a name-to-IP map. The deepest example "
    "reaches into each device's custom_fields to build a software lifecycle report — comparing "
    "the running version against the vendor-recommended one and flagging devices past their "
    "end-of-support date."
)
code(
"# Follow pagination until NetBox stops returning a 'next' URL\n"
"devices = []\n"
"url = f'{NETBOX_URL}/api/dcim/devices/?limit=50'\n"
"while url:\n"
"    body = requests.get(url, headers=headers, verify=False).json()\n"
"    devices.extend(body['results'])\n"
"    url = body['next']"
)

heading("5.3 Script 3 — NetBox: Create a Network Device", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph(
    "POSTs a new device to NetBox — but first resolves the site, device role, and device type "
    "names to the numeric IDs NetBox actually requires, via filtered GET calls. Also checks "
    "whether the device already exists (via GET) before creating it, since NetBox does not "
    "enforce unique device names server-side — a duplicate POST would otherwise silently "
    "create a second device with the same name."
)
code(
"role_id = lookup_id('dcim/device-roles', name='Network Device')\n"
"site_id = lookup_id('dcim/sites', name='Arena Stadium')\n"
"\n"
"payload = {'name': 'Switch-2', 'role': role_id, 'site': site_id, ...}\n"
"resp = requests.post(f'{NETBOX_URL}/api/dcim/devices/', headers=headers, json=payload)"
)

heading("5.4 Script 4 — SSH: Push Baseline Config with Netmiko", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph(
    "Connects to all three lab devices over SSH using Netmiko and pushes banner, logging, "
    "domain, DNS, and NTP settings — the one method that works against every device in this "
    "lab, since none of them expose a working RESTCONF or NETCONF API (confirmed by testing "
    "directly against the hardware), and Cisco's DevNet Always-On IOS-XE sandbox — previously "
    "used here as a working RESTCONF example — is offline for a full platform rebuild, with "
    "Cisco targeting early 2027 for it to return."
)

heading("5.5 Script 5 — ipctl.io: Real Routing Data (Free, No Auth)", 2, (0x2E, 0x74, 0xB5))
doc.add_paragraph(
    "Queries the ipctl.io public API to look up ASN information and IPv4 prefixes — a "
    "maintained replacement for BGPView, which shut down permanently on 2025-11-26. "
    "Students use Cloudflare (AS13335) as a starting point, then look up their own ISP."
)
code(
"# Look up Cloudflare (no auth required)\n"
"resp = requests.get('https://api.ipctl.io/v1/asn/13335')\n"
"data = resp.json()['data']['asn']\n"
"print(data['name'], data['country_code'])\n"
"\n"
"# Same endpoint also returns the ASN's announced prefixes\n"
"prefixes = resp.json()['data']['prefixes']"
)

heading("5.6 Other Free APIs to Explore", 2, (0x2E, 0x74, 0xB5))
table(
    ["API", "Example Call", "Returns"],
    [
        ["PeeringDB",  "GET /api/ix?name=AMS-IX",     "Internet exchange details"],
        ["ipinfo.io",  "GET https://ipinfo.io/8.8.8.8","IP geo, ASN, org"],
        ["httpbin.org","POST /post  (any JSON body)",   "Echo of your request"],
    ],
    [1.5, 3.0, 2.0]
)

# ── 6. End-to-End Pipeline ────────────────────────────────────────────────────

heading("6. End-to-End Automation Pipeline")
doc.add_paragraph(
    "Combining all three sections into a repeatable, auditable change workflow:"
)
code(
"CHANGE REQUEST\n"
"      │\n"
"  git checkout -b change/CR-XXXX      ← branch from main\n"
"      │\n"
"  Edit YAML vars / Jinja2 template    ← define desired state\n"
"      │\n"
"  ansible-playbook --check --diff     ← preview diff\n"
"      │\n"
"  Pull Request → Peer Review          ← human gate\n"
"      │\n"
"  git merge --no-ff                   ← merge to main\n"
"      │\n"
"  ansible-playbook (apply)            ← push to devices\n"
"      │\n"
"  python3 01_netbox_get_devices.py    ← verify via NetBox API\n"
"      │\n"
"  git tag v1.x                        ← tag golden state"
)

# ── 7. Completion Checklist ───────────────────────────────────────────────────

heading("7. Lab Completion Checklist")
table(
    ["Section", "Criteria", "Done?"],
    [
        ["Git",     "Repo initialised with at least one config file",           "☐"],
        ["Git",     "At least two commits in git log",                          "☐"],
        ["Git",     "Feature branch created and merged via merge request",      "☐"],
        ["Git",     "Config rollback demonstrated with git revert",             "☐"],
        ["Ansible", "Inventory with routers + switches groups",                 "☐"],
        ["Ansible", "gather-facts playbook runs on all devices",                "☐"],
        ["Ansible", "push-config applies change via Jinja2 template",           "☐"],
        ["Ansible", "compliance-check generates a report",                      "☐"],
        ["Ansible", "--check --diff used before every real push",               "☐"],
        ["API",     "Script 1 queries NetBox and adds a device via POST",       "☐"],
        ["API",     "Script 2 pages through NetBox devices and summarizes them", "☐"],
        ["API",     "Script 3 resolves IDs and creates a device in NetBox",     "☐"],
        ["API",     "Script 4 pushes baseline config to all devices via SSH",   "☐"],
        ["API",     "Script 5 looks up an ASN and parses BGP prefixes",         "☐"],
        ["Pipeline","Full CR-to-deploy pipeline demonstrated end-to-end",       "☐"],
    ],
    [1.2, 4.2, 0.8]
)

# ── Footer ────────────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f"Network Automation Bootcamp v2.0  ·  {datetime.date.today().strftime('%B %Y')}  "
    "·  10–12 Students"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

output = "/home/jgo/Documents/Automation/Network-Automation-Bootcamp/Network_Automation_Bootcamp_Lab_Guide.docx"
doc.save(output)
print(f"Saved: {output}")
